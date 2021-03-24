import docx


def read_txt_file(filename: str):
    file = open(filename, 'r')
    data = file.read().strip()
    file.close()
    return data


def write_to_txt_file(text: str, filename: str):
    file = open(filename, 'w')
    file.write(text)
    file.close()
    return text


def read_docx_file(filename: str):
    document = docx.Document(filename)
    text = []
    for para in document.paragraphs:
        text.append(para.text)
    return '\n'.join(text)


def write_to_docx_file(text: str, filename: str):
    document = docx.Document()
    document.add_paragraph(text)
    document.save(filename)
    return text
